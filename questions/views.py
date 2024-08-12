import re
from django.conf import settings
from django.shortcuts import render, redirect
from django.contrib import messages
from django.views.decorators.csrf import csrf_exempt
from django.http import JsonResponse, HttpResponse
import zipfile
import base64
import pandas as pd
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from .forms import QuestaoForm
from django.core.files.base import ContentFile
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from .models import Questao, ImagemEnunciado
import json
from bs4 import BeautifulSoup  
from fpdf import FPDF
import os
import tempfile
from PIL import Image
import io

def add_question(request):
    if request.method == 'POST':
        form = QuestaoForm(request.POST, request.FILES)
        if form.is_valid():
            questao = form.save()
            for i, file in enumerate(request.FILES.getlist('imagens')):
                ImagemEnunciado.objects.create(
                    questao=questao,
                    imagem=file,
                    posicao=request.POST.getlist('posicoes')[i]
                )
            messages.success(request, 'Questão adicionada com sucesso!')
            return redirect('adicionar_questao')
    else:
        form = QuestaoForm()
    return render(request, 'questions/adicionar_questao.html', {'form': form})

def questions_list(request):
    questions = Questao.objects.all()
    for question in questions:
        question.short_enunciado = ' '.join(BeautifulSoup(question.enunciado, "html.parser").text.split()[:10]) + '...'
    return render(request, 'questions/questions_list.html', {'questions': questions})

@csrf_exempt
def salvar_questao(request):
    if request.method == 'POST':
        enunciado = request.POST.get('enunciado')
        disciplina = request.POST.get('disciplina')
        gabarito = request.POST.get('gabarito')
        imagens_json = request.POST.getlist('imagens')

        # Substituir os identificadores pelas imagens corretas
        for img_json in imagens_json:
            img_dict = json.loads(img_json)
            img_tag = f'<img src="{img_dict["data"]}" style="width:150px; display:block; margin:10px 0;">'
            enunciado = enunciado.replace(f'{{{{{img_dict["id"]}}}}}', img_tag)
        
        questao = Questao.objects.create(
            enunciado=enunciado,
            disciplina=disciplina,
            gabarito=gabarito
        )
        questao.save()
        
        # Redireciona o usuário para a lista de questões após salvar
        return JsonResponse({'success': True, 'redirect_url': '/questions/lista_questoes/'})
    else:
        return JsonResponse({'success': False, 'error': 'Método inválido.'})


def gerar_simulados(request):
    if request.method == 'POST':
        questoes = Questao.objects.all()
        if not questoes:
            return HttpResponse("Nenhuma questão adicionada.", status=400)

        cadernos = ['Padrão', 'Azul', 'Amarelo', 'Cinza', 'Rosa']
        df = pd.DataFrame(list(questoes.values()))
        df['Indice'] = df.index + 1
        embaralhados = {'Padrão': df}

        for caderno in cadernos[1:]:
            embaralhados[caderno] = df.sample(frac=1).reset_index(drop=True)

        # Criar um arquivo ZIP em memória
        zip_buffer = BytesIO()
        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:

            # Gerar documentos Word
            for caderno, df_caderno in embaralhados.items():
                doc = Document()
                doc.add_heading(f"Simulado {caderno}", 0)

                section = doc.sections[0]
                section.top_margin = Inches(0.5)
                section.bottom_margin = Inches(0.5)
                section.left_margin = Inches(0.5)
                section.right_margin = Inches(0.5)

                # Configurar as colunas
                sectPr = section._sectPr
                cols = OxmlElement('w:cols')
                cols.set(qn('w:num'), '2')
                cols.set(qn('w:sep'), '1')
                sectPr.append(cols)

                # Calcular a largura máxima para a imagem na coluna (largura da coluna)
                column_width = (section.page_width - section.left_margin - section.right_margin) / 2  # 2 colunas

                for indice, (_, row) in enumerate(df_caderno.iterrows(), start=1):
                    doc.add_paragraph(f"Questão {indice}", style='Heading 2')

                    enunciado_html = row['enunciado']
                    soup = BeautifulSoup(enunciado_html, 'html.parser')

                    # Adicionando partes do texto e imagens na ordem correta
                    for element in soup.children:
                        if element.name == 'img':
                            if element.get('src', '').startswith('data:image/'):
                                img_base64 = re.search(r'base64,(.+)', element['src']).group(1)
                                img_data = base64.b64decode(img_base64)
                                img_io = BytesIO(img_data)

                                # Usar PIL para abrir a imagem e verificar as dimensões
                                img = Image.open(img_io)
                                img_width, img_height = img.size

                                # Verificar se a largura da imagem é maior que a largura máxima permitida na coluna
                                if img_width > column_width:
                                    # Calcular a nova altura para manter a proporção da imagem
                                    new_height = int((column_width / img_width) * img_height)
                                    img = img.resize((int(column_width), new_height), Image.ANTIALIAS)
                                
                                # Salvar a imagem redimensionada em BytesIO
                                img_temp_io = BytesIO()
                                img.save(img_temp_io, format='PNG')
                                img_temp_io.seek(0)

                                # Adicionar a imagem ao documento Word
                                doc.add_picture(img_temp_io, width=Inches(column_width / 914400))  # Convertendo de Twips para Inches
                        else:
                            # Aqui vamos garantir que o texto seja adicionado corretamente
                            text = element if isinstance(element, str) else element.get_text()
                            doc.add_paragraph(text)

                doc_io = BytesIO()
                doc.save(doc_io)
                doc_io.seek(0)

                # Adicionar o documento ao ZIP
                zip_file.writestr(f'Simulado_{caderno}.docx', doc_io.read())

            # Gerar planilha Excel
            correlacao = pd.DataFrame(columns=['Disciplina', 'Gabarito', 'Padrão', 'Azul', 'Amarelo', 'Cinza', 'Rosa'])

            for i, questao in df.iterrows():
                correlacao.loc[i] = [
                    questao['disciplina'],
                    questao['gabarito'],
                    embaralhados['Padrão'].loc[embaralhados['Padrão']['Indice'] == questao['Indice']].index[0] + 1,
                    embaralhados['Azul'].loc[embaralhados['Azul']['Indice'] == questao['Indice']].index[0] + 1,
                    embaralhados['Amarelo'].loc[embaralhados['Amarelo']['Indice'] == questao['Indice']].index[0] + 1,
                    embaralhados['Cinza'].loc[embaralhados['Cinza']['Indice'] == questao['Indice']].index[0] + 1,
                    embaralhados['Rosa'].loc[embaralhados['Rosa']['Indice'] == questao['Indice']].index[0] + 1
                ]

            # Salvar o Excel em um BytesIO usando openpyxl
            excel_io = BytesIO()
            with pd.ExcelWriter(excel_io, engine='openpyxl') as writer:
                correlacao.to_excel(writer, index=False)
            excel_io.seek(0)

            # Adicionar o arquivo Excel ao ZIP
            zip_file.writestr('correlacao_questoes.xlsx', excel_io.read())

        zip_buffer.seek(0)

        response = HttpResponse(zip_buffer, content_type='application/zip')
        response['Content-Disposition'] = 'attachment; filename=simulados.zip'
        return response

    return render(request, 'questoes/gerar_simulados.html')

from django.shortcuts import render, get_object_or_404, redirect
from .models import Questao
from .forms import QuestaoForm
from django.contrib import messages
import html
from django.shortcuts import get_object_or_404, render, redirect
from django.contrib import messages
from .models import Questao
from .forms import QuestaoForm

def edit_question(request, pk):
    questao = get_object_or_404(Questao, pk=pk)
    
    if request.method == 'POST':
        enunciado = request.POST.get('enunciado')
        disciplina = request.POST.get('disciplina')
        gabarito = request.POST.get('gabarito')
        
        if enunciado:
            questao.enunciado = enunciado
        if disciplina:
            questao.disciplina = disciplina
        if gabarito:
            questao.gabarito = gabarito
        
        questao.save()
        
        messages.success(request, 'Questão atualizada com sucesso!')
        return redirect('lista_questoes')
    
    else:
        # Decodificar entidades HTML no enunciado
        questao.enunciado = html.unescape(questao.enunciado)
        form = QuestaoForm(instance=questao)
        
    return render(request, 'questions/editar_questao.html', {'form': form})

def delete_question(request, pk):
    questao = get_object_or_404(Questao, pk=pk)
    if request.method == 'POST':
        questao.delete()
        messages.success(request, 'Questão excluída com sucesso!')
        return redirect('lista_questoes')
    return render(request, 'questions/excluir_questao.html', {'questao': questao})
